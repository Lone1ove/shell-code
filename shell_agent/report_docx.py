from pathlib import Path
import re

from docx import Document
from docx.shared import Pt


_INLINE_CODE_RE = re.compile(r"`([^`]+)`")


def _add_text_with_inline_code(paragraph, text: str) -> None:
    text = str(text or "")
    pos = 0
    for match in _INLINE_CODE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        run = paragraph.add_run(match.group(1))
        run.font.name = "Consolas"
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def convert_markdown_to_docx(markdown_path: str | Path, docx_path: str | Path) -> str:
    markdown_path = Path(markdown_path)
    docx_path = Path(docx_path)
    text = markdown_path.read_text(encoding="utf-8-sig")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)

    current_list_mode = None
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            current_list_mode = None
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Title")
            _add_text_with_inline_code(p, stripped[2:].strip())
            current_list_mode = None
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            _add_text_with_inline_code(p, stripped[3:].strip())
            current_list_mode = None
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            _add_text_with_inline_code(p, stripped[4:].strip())
            current_list_mode = None
            continue

        if stripped.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            _add_text_with_inline_code(p, stripped[5:].strip())
            current_list_mode = None
            continue

        bullet_match = re.match(r"^- (.+)$", stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_text_with_inline_code(p, bullet_match.group(1))
            current_list_mode = "bullet"
            continue

        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if number_match:
            p = doc.add_paragraph(style="List Number")
            _add_text_with_inline_code(p, number_match.group(1))
            current_list_mode = "number"
            continue

        p = doc.add_paragraph()
        _add_text_with_inline_code(p, stripped)
        current_list_mode = None

    doc.save(docx_path)
    return str(docx_path)
